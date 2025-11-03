"""
DOCX Export Service
Converts chapters to Microsoft Word documents (.docx)
"""

import logging
import re
from typing import Dict, List, Optional, Any
from datetime import datetime
from io import BytesIO

# Make python-docx optional (might not be installed in all environments)
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    DOCX_AVAILABLE = True
except ImportError as e:
    DOCX_AVAILABLE = False
    logger_temp = logging.getLogger(__name__)
    logger_temp.warning(f"python-docx not available: {e}. DOCX export will not work.")

import markdown2

from backend.database.models.chapter import Chapter

logger = logging.getLogger(__name__)


class DOCXExporter:
    """
    Export chapters to Microsoft Word DOCX format
    """

    def __init__(self):
        self.document = None

    def export_chapter_to_docx(
        self,
        chapter: Chapter,
        include_images: bool = True,
        include_quality_metrics: bool = True
    ) -> bytes:
        """
        Export chapter to DOCX format

        Args:
            chapter: Chapter object to export
            include_images: Whether to include images
            include_quality_metrics: Include quality and confidence metrics

        Returns:
            DOCX file as bytes

        Raises:
            RuntimeError: If python-docx is not available
        """
        if not DOCX_AVAILABLE:
            raise RuntimeError(
                "python-docx library is not available. Please install it with: pip install python-docx"
            )

        logger.info(f"Exporting chapter {chapter.id} to DOCX")

        # Create new document
        self.document = Document()

        # Set up styles
        self._setup_styles()

        # Add title page
        self._add_title_page(chapter)

        # Add quality metrics if requested
        if include_quality_metrics:
            self._add_quality_metrics(chapter)

        # Add page break
        self.document.add_page_break()

        # Add sections
        self._add_sections(chapter.sections or [], include_images)

        # Add references
        self._add_references(chapter)

        # Save to BytesIO
        docx_bytes = BytesIO()
        self.document.save(docx_bytes)
        docx_bytes.seek(0)

        logger.info(f"DOCX generated successfully for chapter {chapter.id}")

        return docx_bytes.getvalue()

    def _setup_styles(self):
        """Set up custom document styles"""
        styles = self.document.styles

        # Title style
        if 'Chapter Title' not in styles:
            title_style = styles.add_style('Chapter Title', WD_STYLE_TYPE.PARAGRAPH)
            title_font = title_style.font
            title_font.name = 'Arial'
            title_font.size = Pt(24)
            title_font.bold = True
            title_font.color.rgb = RGBColor(0, 51, 102)  # Dark blue
            title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_style.paragraph_format.space_after = Pt(12)

        # Subtitle style
        if 'Subtitle' not in styles:
            subtitle_style = styles.add_style('Subtitle', WD_STYLE_TYPE.PARAGRAPH)
            subtitle_font = subtitle_style.font
            subtitle_font.name = 'Arial'
            subtitle_font.size = Pt(14)
            subtitle_font.color.rgb = RGBColor(64, 64, 64)
            subtitle_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            subtitle_style.paragraph_format.space_after = Pt(6)

        # Section Heading style
        if 'Section Heading' not in styles:
            section_style = styles.add_style('Section Heading', WD_STYLE_TYPE.PARAGRAPH)
            section_font = section_style.font
            section_font.name = 'Arial'
            section_font.size = Pt(16)
            section_font.bold = True
            section_font.color.rgb = RGBColor(0, 51, 102)
            section_style.paragraph_format.space_before = Pt(12)
            section_style.paragraph_format.space_after = Pt(6)

    def _add_title_page(self, chapter: Chapter):
        """Add title page to document"""
        # Main title
        title = self.document.add_paragraph(chapter.title or 'Untitled', style='Chapter Title')

        # Subtitle
        subtitle = self.document.add_paragraph('Neurosurgical Core of Knowledge', style='Subtitle')

        # Add spacing
        self.document.add_paragraph()

        # Metadata table
        table = self.document.add_table(rows=5, cols=2)
        table.style = 'Light Grid Accent 1'

        # Version
        row = table.rows[0]
        row.cells[0].text = 'Version'
        row.cells[1].text = str(chapter.version or '1.0')

        # Generated date
        row = table.rows[1]
        row.cells[0].text = 'Generated'
        row.cells[1].text = datetime.now().strftime('%B %d, %Y')

        # Chapter type
        row = table.rows[2]
        row.cells[0].text = 'Chapter Type'
        row.cells[1].text = chapter.chapter_type or 'N/A'

        # Word count
        row = table.rows[3]
        row.cells[0].text = 'Total Words'
        row.cells[1].text = str(chapter.total_words or chapter.word_count or 0)

        # Sections
        row = table.rows[4]
        row.cells[0].text = 'Sections'
        row.cells[1].text = str(chapter.total_sections or len(chapter.sections or []))

        # Add spacing
        self.document.add_paragraph()

        # Summary/Abstract
        if chapter.summary:
            summary_heading = self.document.add_paragraph('Summary')
            summary_heading.runs[0].bold = True
            summary_heading.runs[0].font.size = Pt(14)

            summary_para = self.document.add_paragraph(chapter.summary)
            summary_para.paragraph_format.space_after = Pt(12)

    def _add_quality_metrics(self, chapter: Chapter):
        """Add quality metrics section"""
        # Add page break
        self.document.add_page_break()

        # Quality metrics heading
        heading = self.document.add_paragraph('Quality & Confidence Metrics')
        heading.runs[0].bold = True
        heading.runs[0].font.size = Pt(16)
        heading.runs[0].font.color.rgb = RGBColor(0, 51, 102)

        quality_scores = chapter.quality_scores or {}
        gen_conf = chapter.generation_confidence_data or {}

        # Overall quality score
        if quality_scores:
            overall_para = self.document.add_paragraph()
            overall_para.add_run('Overall Quality: ').bold = True
            overall_text = f"{quality_scores.get('overall', 0) * 100:.1f}% ({quality_scores.get('rating', 'N/A')})"
            overall_para.add_run(overall_text)

        # Generation confidence
        if gen_conf:
            conf_para = self.document.add_paragraph()
            conf_para.add_run('Generation Confidence: ').bold = True
            conf_text = f"{gen_conf.get('overall', 0) * 100:.1f}% ({gen_conf.get('rating', 'N/A')})"
            conf_para.add_run(conf_text)

        # Detailed scores table
        if quality_scores:
            self.document.add_paragraph()
            detail_heading = self.document.add_paragraph('Detailed Quality Scores')
            detail_heading.runs[0].bold = True
            detail_heading.runs[0].font.size = Pt(12)

            table = self.document.add_table(rows=5, cols=2)
            table.style = 'Light List Accent 1'

            metrics = [
                ('Depth', quality_scores.get('depth', 0)),
                ('Coverage', quality_scores.get('coverage', 0)),
                ('Currency', quality_scores.get('currency', 0)),
                ('Evidence', quality_scores.get('evidence', 0))
            ]

            for i, (metric_name, score) in enumerate(metrics):
                row = table.rows[i]
                row.cells[0].text = metric_name
                row.cells[1].text = f"{score * 100:.1f}%"

        # Confidence breakdown
        if gen_conf and 'breakdown' in gen_conf:
            self.document.add_paragraph()
            breakdown_heading = self.document.add_paragraph('Confidence Breakdown')
            breakdown_heading.runs[0].bold = True
            breakdown_heading.runs[0].font.size = Pt(12)

            components = gen_conf['breakdown'].get('components', {})

            for component_name, component_data in components.items():
                if isinstance(component_data, dict):
                    score = component_data.get('score', 0)
                    weight = component_data.get('weight', 0)
                    description = component_data.get('description', '')

                    para = self.document.add_paragraph()
                    para.add_run(f"{component_name.replace('_', ' ').title()}: ").bold = True
                    para.add_run(f"{score * 100:.1f}% (Weight: {weight * 100:.0f}%)")
                    if description:
                        desc_para = self.document.add_paragraph(f"   {description}")
                        desc_para.paragraph_format.left_indent = Inches(0.25)

    def _add_sections(self, sections: List[Dict[str, Any]], include_images: bool):
        """Add chapter sections to document"""
        for i, section in enumerate(sections, 1):
            title = section.get('title', f'Section {i}')
            content = section.get('content', '')

            # Section heading
            heading = self.document.add_paragraph(title, style='Section Heading')

            # Convert HTML content to plain text with some formatting
            if content:
                self._add_html_content(content, include_images)

            # Add spacing between sections
            self.document.add_paragraph()

    def _add_html_content(self, html_content: str, include_images: bool):
        """
        Convert HTML content to Word paragraphs

        Args:
            html_content: HTML formatted content
            include_images: Whether to include images
        """
        # Simple HTML to text conversion
        # Strip HTML tags but preserve some formatting

        # Remove script and style elements
        html_content = re.sub(r'<script[^>]*?>.*?</script>', '', html_content, flags=re.DOTALL | re.IGNORECASE)
        html_content = re.sub(r'<style[^>]*?>.*?</style>', '', html_content, flags=re.DOTALL | re.IGNORECASE)

        # Handle headings
        html_content = re.sub(r'<h3[^>]*?>(.*?)</h3>', r'\n\n**\1**\n', html_content, flags=re.DOTALL | re.IGNORECASE)
        html_content = re.sub(r'<h4[^>]*?>(.*?)</h4>', r'\n\n*\1*\n', html_content, flags=re.DOTALL | re.IGNORECASE)

        # Handle lists
        html_content = re.sub(r'<li[^>]*?>(.*?)</li>', r'• \1\n', html_content, flags=re.DOTALL | re.IGNORECASE)
        html_content = re.sub(r'</?[ou]l[^>]*?>', '\n', html_content, flags=re.IGNORECASE)

        # Handle paragraphs
        html_content = re.sub(r'<p[^>]*?>(.*?)</p>', r'\1\n\n', html_content, flags=re.DOTALL | re.IGNORECASE)

        # Handle bold and italic
        html_content = re.sub(r'<strong[^>]*?>(.*?)</strong>', r'**\1**', html_content, flags=re.DOTALL | re.IGNORECASE)
        html_content = re.sub(r'<b[^>]*?>(.*?)</b>', r'**\1**', html_content, flags=re.DOTALL | re.IGNORECASE)
        html_content = re.sub(r'<em[^>]*?>(.*?)</em>', r'*\1*', html_content, flags=re.DOTALL | re.IGNORECASE)
        html_content = re.sub(r'<i[^>]*?>(.*?)</i>', r'*\1*', html_content, flags=re.DOTALL | re.IGNORECASE)

        # Handle line breaks
        html_content = re.sub(r'<br\s*/?>', '\n', html_content, flags=re.IGNORECASE)

        # Remove remaining HTML tags
        html_content = re.sub(r'<[^>]+?>', '', html_content)

        # Decode HTML entities
        html_content = html_content.replace('&nbsp;', ' ')
        html_content = html_content.replace('&amp;', '&')
        html_content = html_content.replace('&lt;', '<')
        html_content = html_content.replace('&gt;', '>')
        html_content = html_content.replace('&quot;', '"')

        # Split into paragraphs and add to document
        paragraphs = html_content.split('\n\n')

        for para_text in paragraphs:
            para_text = para_text.strip()
            if not para_text:
                continue

            para = self.document.add_paragraph()

            # Handle bold and italic formatting
            parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', para_text)

            for part in parts:
                if not part:
                    continue

                if part.startswith('**') and part.endswith('**'):
                    # Bold text
                    run = para.add_run(part[2:-2])
                    run.bold = True
                elif part.startswith('*') and part.endswith('*'):
                    # Italic text
                    run = para.add_run(part[1:-1])
                    run.italic = True
                else:
                    # Normal text
                    para.add_run(part)

    def _add_references(self, chapter: Chapter):
        """Add references section"""
        if not chapter.citations:
            return

        # Add page break
        self.document.add_page_break()

        # References heading
        heading = self.document.add_paragraph('References')
        heading.runs[0].bold = True
        heading.runs[0].font.size = Pt(16)
        heading.runs[0].font.color.rgb = RGBColor(0, 51, 102)

        citations = chapter.citations if isinstance(chapter.citations, list) else []

        for i, citation in enumerate(citations, 1):
            if isinstance(citation, dict):
                ref_para = self.document.add_paragraph(style='List Number')

                authors = citation.get('authors', 'Unknown')
                title = citation.get('title', 'Untitled')
                year = citation.get('year', 'n.d.')
                journal = citation.get('journal', '')

                ref_text = f"{authors}. {title}. "
                if journal:
                    ref_text += f"{journal}. "
                ref_text += f"{year}."

                ref_para.add_run(ref_text)

    def export_citations_only(
        self,
        chapter: Chapter,
        citation_style: str = "apa"
    ) -> bytes:
        """
        Export only citations to DOCX

        Args:
            chapter: Chapter object
            citation_style: Citation style (apa, vancouver, chicago)

        Returns:
            DOCX file as bytes
        """
        logger.info(f"Exporting citations for chapter {chapter.id} (style={citation_style})")

        self.document = Document()

        # Title
        title = self.document.add_paragraph(f"References - {chapter.title}")
        title.runs[0].bold = True
        title.runs[0].font.size = Pt(16)

        # Add citations
        self._add_references(chapter)

        # Save to BytesIO
        docx_bytes = BytesIO()
        self.document.save(docx_bytes)
        docx_bytes.seek(0)

        return docx_bytes.getvalue()
