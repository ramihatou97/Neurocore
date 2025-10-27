"""
Export Service - Chapter Export to Multiple Formats
Handles PDF, DOCX, and HTML export with custom templates
"""

from typing import Dict, Any, Optional
from sqlalchemy.orm import Session
from jinja2 import Template
from datetime import datetime
import io
import os

from backend.database.models import Chapter, ExportTemplate, ExportHistory
from backend.services.citation_service import CitationService
from backend.services.bibliography_service import BibliographyService
from backend.utils import get_logger

logger = get_logger(__name__)


class ExportService:
    """
    Main export service for all formats

    Coordinates export process:
    1. Prepare chapter content
    2. Generate bibliography
    3. Apply template
    4. Convert to target format
    5. Log export history
    """

    def __init__(self, db: Session):
        """
        Initialize export service

        Args:
            db: Database session
        """
        self.db = db
        self.citation_service = CitationService(db)
        self.bibliography_service = BibliographyService(db)

    def export_chapter(
        self,
        chapter_id: str,
        user_id: str,
        export_format: str,
        template_id: Optional[str] = None,
        citation_style: str = "apa",
        options: Optional[Dict[str, Any]] = None
    ) -> Dict[str, Any]:
        """
        Export chapter to specified format

        Args:
            chapter_id: Chapter ID to export
            user_id: User performing export
            export_format: Target format (pdf, docx, html)
            template_id: Template ID (optional, uses default if not specified)
            citation_style: Citation style for bibliography
            options: Additional export options

        Returns:
            Dictionary with export results and file info
        """
        try:
            # Get chapter
            chapter = self.db.query(Chapter).filter(Chapter.id == chapter_id).first()
            if not chapter:
                raise ValueError(f"Chapter {chapter_id} not found")

            # Get template
            if template_id:
                template = self.db.query(ExportTemplate).filter(
                    ExportTemplate.id == template_id
                ).first()
            else:
                template = self._get_default_template(export_format)

            if not template:
                raise ValueError(f"No template found for format {export_format}")

            # Prepare content
            content_data = self._prepare_chapter_content(chapter, citation_style, options or {})

            # Generate export based on format
            if export_format == "pdf":
                result = self._export_to_pdf(chapter, template, content_data, options or {})
            elif export_format == "docx":
                result = self._export_to_docx(chapter, template, content_data, options or {})
            elif export_format == "html":
                result = self._export_to_html(chapter, template, content_data, options or {})
            else:
                raise ValueError(f"Unsupported export format: {export_format}")

            # Log export history
            self._log_export(
                chapter_id=chapter_id,
                user_id=user_id,
                template_id=template_id,
                export_format=export_format,
                file_name=result.get("file_name"),
                file_size=result.get("file_size"),
                options=options or {}
            )

            return {
                "success": True,
                "format": export_format,
                "file_name": result.get("file_name"),
                "file_size": result.get("file_size"),
                "content": result.get("content"),
                "message": f"Successfully exported chapter to {export_format.upper()}"
            }

        except Exception as e:
            logger.error(f"Export failed: {str(e)}", exc_info=True)
            return {
                "success": False,
                "error": str(e),
                "message": f"Failed to export chapter: {str(e)}"
            }

    def _prepare_chapter_content(
        self,
        chapter: Chapter,
        citation_style: str,
        options: Dict[str, Any]
    ) -> Dict[str, Any]:
        """
        Prepare chapter content for export

        Args:
            chapter: Chapter to export
            citation_style: Citation style for bibliography
            options: Export options

        Returns:
            Dictionary with prepared content and metadata
        """
        try:
            # Extract content from chapter
            content_html = ""
            if hasattr(chapter, 'sections') and chapter.sections:
                if isinstance(chapter.sections, list):
                    sections_html = []
                    for section in chapter.sections:
                        section_title = section.get("title", "")
                        section_content = section.get("content", "")
                        sections_html.append(f"<h2>{section_title}</h2>\n{section_content}")
                    content_html = "\n\n".join(sections_html)
                else:
                    content_html = str(chapter.sections)
            else:
                content_html = getattr(chapter, 'content', '')

            # Generate bibliography if requested
            bibliography = []
            if options.get("include_bibliography", True):
                bib_result = self.bibliography_service.generate_bibliography(
                    chapter,
                    citation_style,
                    sort_by=options.get("bibliography_sort", "author")
                )
                bibliography = bib_result.get("bibliography", [])

            # Get author information
            author_name = "Unknown Author"
            if hasattr(chapter, 'author') and chapter.author:
                author_name = getattr(chapter.author, 'name', getattr(chapter.author, 'email', 'Unknown'))

            return {
                "title": chapter.title,
                "author_name": author_name,
                "date": datetime.now().strftime("%B %d, %Y"),
                "content": content_html,
                "bibliography": bibliography,
                "chapter_type": getattr(chapter, 'chapter_type', ''),
                "word_count": chapter.get_word_count() if hasattr(chapter, 'get_word_count') else 0,
                "metadata": {
                    "created_at": chapter.created_at.isoformat() if hasattr(chapter, 'created_at') else None,
                    "updated_at": chapter.updated_at.isoformat() if hasattr(chapter, 'updated_at') else None,
                }
            }

        except Exception as e:
            logger.error(f"Failed to prepare chapter content: {str(e)}")
            raise

    def _export_to_html(
        self,
        chapter: Chapter,
        template: ExportTemplate,
        content_data: Dict[str, Any],
        options: Dict[str, Any]
    ) -> Dict[str, Any]:
        """Export to HTML format"""
        try:
            # Render template with Jinja2
            jinja_template = Template(template.template_content)

            # Merge template styles with options
            template_vars = {
                **content_data,
                **template.styles,
                **options
            }

            html_content = jinja_template.render(**template_vars)

            # Generate file name
            file_name = f"{chapter.title.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d')}.html"

            return {
                "file_name": file_name,
                "content": html_content,
                "file_size": f"{len(html_content.encode('utf-8')) / 1024:.2f} KB"
            }

        except Exception as e:
            logger.error(f"HTML export failed: {str(e)}")
            raise

    def _export_to_pdf(
        self,
        chapter: Chapter,
        template: ExportTemplate,
        content_data: Dict[str, Any],
        options: Dict[str, Any]
    ) -> Dict[str, Any]:
        """
        Export to PDF format using WeasyPrint

        Note: This requires WeasyPrint to be installed
        """
        try:
            # First generate HTML
            html_result = self._export_to_html(chapter, template, content_data, options)
            html_content = html_result["content"]

            # Try to import WeasyPrint
            try:
                from weasyprint import HTML
            except ImportError:
                logger.warning("WeasyPrint not installed, returning HTML instead of PDF")
                return {
                    "file_name": html_result["file_name"].replace(".html", ".pdf"),
                    "content": html_content,
                    "file_size": html_result["file_size"],
                    "note": "WeasyPrint not installed - returning HTML content"
                }

            # Convert HTML to PDF
            pdf_bytes = HTML(string=html_content).write_pdf()

            file_name = f"{chapter.title.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d')}.pdf"

            return {
                "file_name": file_name,
                "content": pdf_bytes,
                "file_size": f"{len(pdf_bytes) / 1024:.2f} KB"
            }

        except Exception as e:
            logger.error(f"PDF export failed: {str(e)}")
            raise

    def _export_to_docx(
        self,
        chapter: Chapter,
        template: ExportTemplate,
        content_data: Dict[str, Any],
        options: Dict[str, Any]
    ) -> Dict[str, Any]:
        """
        Export to DOCX format using python-docx

        Note: This requires python-docx to be installed
        """
        try:
            # Try to import python-docx
            try:
                from docx import Document
                from docx.shared import Pt, Inches
                from docx.enum.text import WD_ALIGN_PARAGRAPH
            except ImportError:
                logger.warning("python-docx not installed, returning HTML instead")
                html_result = self._export_to_html(chapter, template, content_data, options)
                return {
                    "file_name": html_result["file_name"].replace(".html", ".docx"),
                    "content": html_result["content"],
                    "file_size": html_result["file_size"],
                    "note": "python-docx not installed - returning HTML content"
                }

            # Create document
            doc = Document()

            # Apply styles from template
            styles = template.styles
            font_family = styles.get("font_family", "Times New Roman")
            font_size = styles.get("font_size", "12pt")

            # Add title
            title_para = doc.add_heading(content_data["title"], level=0)
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Add metadata
            metadata_para = doc.add_paragraph()
            metadata_para.add_run(f"{content_data['author_name']}\n").bold = True
            metadata_para.add_run(content_data["date"])
            metadata_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Add spacing
            doc.add_paragraph()

            # Add content (simplified - in production would parse HTML properly)
            # For now, add as plain text
            content_text = content_data["content"]
            # Strip HTML tags (basic)
            import re
            content_text = re.sub('<.*?>', '', content_text)

            content_para = doc.add_paragraph(content_text)

            # Add bibliography if present
            if content_data.get("bibliography"):
                doc.add_page_break()
                doc.add_heading("References", level=1)

                for entry in content_data["bibliography"]:
                    bib_para = doc.add_paragraph(entry.get("formatted", ""))
                    # Hanging indent
                    bib_para.paragraph_format.left_indent = Inches(0.5)
                    bib_para.paragraph_format.first_line_indent = Inches(-0.5)

            # Save to bytes
            docx_bytes = io.BytesIO()
            doc.save(docx_bytes)
            docx_bytes.seek(0)

            file_name = f"{chapter.title.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d')}.docx"

            return {
                "file_name": file_name,
                "content": docx_bytes.getvalue(),
                "file_size": f"{len(docx_bytes.getvalue()) / 1024:.2f} KB"
            }

        except Exception as e:
            logger.error(f"DOCX export failed: {str(e)}")
            raise

    def _get_default_template(self, export_format: str) -> Optional[ExportTemplate]:
        """Get default template for format"""
        return self.db.query(ExportTemplate).filter(
            ExportTemplate.format == export_format,
            ExportTemplate.is_default == True,
            ExportTemplate.is_public == True
        ).first()

    def _log_export(
        self,
        chapter_id: str,
        user_id: str,
        template_id: Optional[str],
        export_format: str,
        file_name: str,
        file_size: str,
        options: Dict[str, Any]
    ):
        """Log export to history"""
        try:
            export_log = ExportHistory(
                chapter_id=chapter_id,
                user_id=user_id,
                template_id=template_id,
                export_format=export_format,
                file_name=file_name,
                file_size=file_size,
                export_options=options,
                status="completed",
                completed_at=datetime.utcnow()
            )

            self.db.add(export_log)
            self.db.commit()

        except Exception as e:
            logger.error(f"Failed to log export: {str(e)}")
            # Don't fail export if logging fails
            self.db.rollback()
